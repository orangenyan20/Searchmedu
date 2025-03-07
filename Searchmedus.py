import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import time
import io
from docx import Document
from docx.shared import Inches
from PIL import Image

# 画像ダウンロード用関数（Streamlitで表示せずに画像をワードファイルに保存）
def download_image(url):
    response = requests.get(url)
    if response.status_code == 200:
        img = Image.open(io.BytesIO(response.content))
        return img
    return None

# 検索してURLを取得する関数
def search_and_scrape(search_query):
    search_query = search_query.strip().replace(' ', '%20')
    result_links = []
    page_num = 1  # 1ページ目から開始
    pattern = re.compile(r'/([1-9][0-9]{2,})[A-Za-z]\d{2}')
    
    while page_num <= 6:
        if page_num == 1:
            url = f'https://medu4.com/quizzes/result?q={search_query}&st=all'
        else:
            url = f'https://medu4.com/quizzes/result?page={page_num}&q={search_query}&st=all'

        response = requests.get(url)
        if response.status_code != 200:
            break  

        soup = BeautifulSoup(response.text, 'html.parser')
        all_links = [link['href'] for link in soup.find_all('a', href=True)]
        page_results = [link for link in all_links if pattern.search(link)]

        if not page_results:
            break

        result_links.extend(page_results)
        page_num += 1  
        time.sleep(0.5)

    full_urls = [f"https://medu4.com{link}" for link in result_links]
    return full_urls

# ページ内容を取得する関数
def get_page_text(url):
    response = requests.get(url)
    soup = BeautifulSoup(response.text, 'html.parser')

    category = soup.find('span', class_='button-small-line')
    category_name = category.text.strip() if category else '分野名なし'

    problem = soup.find('div', class_='quiz-body mb-64')
    problem_text = problem.text.strip() if problem else '問題文なし'

    choices = []
    for choice in soup.find_all('div', class_='box-select'):
        choice_header = choice.find('span', {'class': 'choice-header'}).text.strip()
        choice_text = choice.find_all('span')[1].text.strip()  
        choices.append(f"{choice_header} {choice_text}")

    answer = soup.find('h4')
    answer_text = answer.text.strip() if answer else '解答なし'

    explanation = soup.find('div', class_='explanation')
    explanation_text = explanation.text.strip() if explanation else '解説なし'

    # 画像URLを取得（指定された部分のみ）
    image_urls = []
    image_divs = soup.find_all('div', class_='box-quiz-image mb-32')
    for div in image_divs:
        img_tag = div.find('img')
        if img_tag and img_tag.get('src'):
            img_url = img_tag['src']
            # サムネイルのURLから元画像のURLを取得
            img_url_full = img_url.replace('thumb_', '')  # サムネイル画像からオリジナル画像URLに変換
            image_urls.append(img_url_full)

    return {
        "category": category_name,
        "problem": problem_text,
        "choices": choices,
        "answer": answer_text,
        "explanation": explanation_text,
        "images": image_urls
    }

# Wordファイルを作成する関数
def create_word_doc(pages_data, search_query):
    doc = Document()
    doc.add_heading('検索結果', 0)

    for page_data in pages_data:
        doc.add_paragraph(f"問題文: {page_data['problem']}")
        
        # 画像があれば、問題文の真下に画像を追加（表示はせず、Word内のみ）
        if page_data['images']:
            for img_url in page_data['images']:
                img = download_image(img_url)
                if img:
                    # 画像をWordファイルに追加
                    with io.BytesIO() as img_stream:
                        img.save(img_stream, format='PNG')
                        img_stream.seek(0)
                        doc.add_picture(img_stream, width=Inches(3))  # 適切なサイズに調整
                else:
                    doc.add_paragraph(f"画像取得失敗: {img_url}")

        doc.add_paragraph("選択肢:")
        for choice in page_data['choices']:
            doc.add_paragraph(choice)
        doc.add_paragraph(f"解答: {page_data['answer']}")
        doc.add_paragraph(f"解説: {page_data['explanation']}")

        doc.add_paragraph("-" * 50)

    filename = f"{search_query}_search_results.docx"
    doc.save(filename)
    return filename

# Streamlit UI
st.title("Medu4 検索ツール")
search_query = st.text_input("検索ワードを入力してください")

if st.button("検索"):
    with st.spinner("検索中..."):
        result_pages = search_and_scrape(search_query)

    if result_pages:
        pages_data = [get_page_text(url) for url in result_pages]

        with st.spinner("ワードファイル作成中..."):
            filename = create_word_doc(pages_data, search_query)
        
        st.success("検索結果を Word に保存しました！")
        with open(filename, "rb") as file:
            st.download_button("📄 Wordファイルをダウンロード", file, file_name=filename)
    else:
        st.error("検索結果がありませんでした。")
